#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AHP 플랫폼 분석 보고서 생성 스크립트
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_color(doc, text, level, color=None):
    """색상이 있는 제목 추가"""
    heading = doc.add_heading(text, level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_paragraph_with_style(doc, text, bold=False, italic=False, color=None, size=None):
    """스타일이 적용된 문단 추가"""
    para = doc.add_paragraph(text)
    if bold or italic or color or size:
        run = para.runs[0]
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return para

def add_table_with_style(doc, data, has_header=True):
    """스타일이 적용된 표 추가"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            
            # 헤더 행 스타일링
            if has_header and i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # 배경색 설정
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2563EB')  # 파란색
                cell._element.get_or_add_tcPr().append(shading_elm)
    
    return table

def create_analysis_report():
    """분석 보고서 생성"""
    doc = Document()
    
    # 문서 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ==================== 표지 ====================
    title = doc.add_heading('AHP 의사결정 지원 플랫폼', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)  # 파란색
    
    subtitle = doc.add_heading('심층 분석 및 평가 보고서', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(107, 114, 128)  # 회색
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 보고서 정보
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}\n')
    info_run.font.size = Pt(12)
    info_run = info_para.add_run('버전: 2.1.1\n')
    info_run.font.size = Pt(12)
    info_run = info_para.add_run('분석 도구: Claude AI & Python')
    info_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ==================== 목차 ====================
    add_heading_with_color(doc, '목차', 1, RGBColor(37, 99, 235))
    
    toc_items = [
        '1. 프로젝트 개요',
        '2. 종합 평가',
        '3. 치명적 문제 (Critical Issues)',
        '4. 중요 문제 (High Priority Issues)',
        '5. 경미한 문제 (Medium Priority Issues)',
        '6. 성능 문제 분석',
        '7. 아키텍처 문제',
        '8. 데이터베이스 및 백엔드 문제',
        '9. 개발 환경 문제',
        '10. 우선순위별 해결 로드맵',
        '11. 긍정적 측면',
        '12. 최종 권고사항',
        '13. 결론'
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Number')
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ==================== 1. 프로젝트 개요 ====================
    add_heading_with_color(doc, '1. 프로젝트 개요', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('1.1 프로젝트 정보', 2)
    
    project_data = [
        ['항목', '내용'],
        ['프로젝트명', 'AHP for Paper - Multi-Layer Decision Support Platform'],
        ['목적', '연구자를 위한 전문적인 다기준 의사결정 지원 시스템'],
        ['완성도', '75% (MVP 목표: 2024년 12월 31일)'],
        ['코드 규모', '26,571 라인 (247개 컴포넌트)'],
        ['배포 환경', 'GitHub Pages (Frontend) + Render.com (Backend)'],
        ['최근 업데이트', '2024-11-11 - 계층적 평가 시스템 완전 구현']
    ]
    add_table_with_style(doc, project_data)
    
    doc.add_paragraph('')
    doc.add_heading('1.2 기술 스택', 2)
    
    tech_data = [
        ['구분', '기술', '버전'],
        ['Frontend', 'React + TypeScript', '18.2 + 4.9.5'],
        ['', 'Tailwind CSS', '3.4.17'],
        ['', 'Recharts (시각화)', '2.12.7'],
        ['Backend', 'Django REST Framework', '5.0'],
        ['', 'PostgreSQL', '17.0'],
        ['DevOps', 'GitHub Actions', 'CI/CD'],
        ['', 'GitHub Pages', '프론트엔드 호스팅'],
        ['', 'Render.com', '백엔드 호스팅 (유료)']
    ]
    add_table_with_style(doc, tech_data)
    
    doc.add_paragraph('')
    doc.add_heading('1.3 주요 기능', 2)
    
    features = [
        '계층적 AHP 모델링 및 쌍대비교 평가',
        '다중 평가자 협업 시스템 (실시간 초대 및 권한 관리)',
        'Power Method 알고리즘 기반 고유벡터 계산',
        'CR(Consistency Ratio) 자동 검증',
        '민감도 분석 및 What-if 시나리오 시뮬레이션',
        'AI 통합 기능 (챗봇, 해석, 자료 생성)',
        '역할 기반 접근 제어 (Super Admin, Researcher, Evaluator)',
        '데이터 시각화 대시보드',
        'QR 코드 기반 익명 평가자 할당',
        'Excel/PDF 결과 내보내기'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 2. 종합 평가 ====================
    add_heading_with_color(doc, '2. 종합 평가', 1, RGBColor(37, 99, 235))
    
    para = doc.add_paragraph()
    run = para.add_run('심각도 분류: ')
    run.font.bold = True
    run = para.add_run('🔴 치명적 | 🟠 중요 | 🟡 경미')
    
    doc.add_paragraph('')
    
    score_data = [
        ['평가 항목', '점수', '등급', '평가'],
        ['기능 완성도', '75/100', '🟢', '양호'],
        ['코드 품질', '55/100', '🟡', '개선 필요'],
        ['보안', '40/100', '🔴', '위험'],
        ['성능', '45/100', '🔴', '위험'],
        ['테스트 커버리지', '30/100', '🔴', '치명적'],
        ['문서화', '85/100', '🟢', '우수'],
        ['전체 평균', '55/100', '🟡', '개선 필요']
    ]
    add_table_with_style(doc, score_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '주요 발견 사항:', bold=True, size=12)
    
    findings = [
        '보안 취약점 다수 발견 (React Router XSS, ajv ReDoS)',
        '번들 크기 과다 (2.8MB) - 초기 로딩 5초 이상 예상',
        'PersonalServiceDashboard.tsx 5,345 라인 - 유지보수 불가능',
        '디버그 코드 568개 (console.log) 프로덕션 포함',
        '테스트 커버리지 10% 미만 (20개 테스트 vs 247개 컴포넌트)',
        '패키지 업데이트 지연 (TypeScript 4.9 → 5.9, React 18 → 19)',
        'Code Splitting 미구현으로 초기 로딩 성능 저하'
    ]
    
    for finding in findings:
        para = doc.add_paragraph(finding, style='List Bullet')
        if '보안' in finding or '2.8MB' in finding or '5,345' in finding:
            for run in para.runs:
                run.font.color.rgb = RGBColor(220, 38, 38)  # 빨간색
    
    doc.add_page_break()
    
    # ==================== 3. 치명적 문제 ====================
    add_heading_with_color(doc, '3. 치명적 문제 (Critical Issues)', 1, RGBColor(220, 38, 38))
    
    doc.add_heading('3.1 보안 취약점 - 즉시 조치 필요 🔴🔴🔴', 2)
    
    para = doc.add_paragraph()
    run = para.add_run('보안 감사 결과:\n')
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    security_issues = [
        'React Router XSS 취약점 (GHSA-2w69-qvjg-hvjx) - HIGH Severity',
        'ajv ReDoS 취약점 (GHSA-2g4f-4pwh-qvx6) - MODERATE Severity',
        '다수의 의존성 체인 취약점'
    ]
    
    for issue in security_issues:
        doc.add_paragraph(issue, style='List Bullet')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '영향 범위: 프로덕션 배포 전체', bold=True, color=RGBColor(220, 38, 38))
    add_paragraph_with_style(doc, '해결 방법:', bold=True)
    doc.add_paragraph('npm audit fix', style='List Bullet')
    doc.add_paragraph('npm audit fix --force (주의: 일부 breaking changes 발생 가능)', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('3.2 번들 크기 과다 - 성능 위기 🔴🔴', 2)
    
    bundle_data = [
        ['항목', '크기', '상태'],
        ['main.js (압축 전)', '2.8MB', '⚠️ 위험'],
        ['node_modules', '642MB', '과다'],
        ['의존성 패키지 수', '1,042개', '정리 필요'],
        ['빌드 결과물 (전체)', '12MB', '최적화 필요']
    ]
    add_table_with_style(doc, bundle_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True)
    problems = [
        '초기 로딩 시간 5초 이상 예상',
        '모바일 환경에서 치명적 지연',
        '대역폭 낭비 심각',
        'Code Splitting 미구현',
        'Lazy Loading 미적용',
        'Tree Shaking 최적화 부족'
    ]
    for problem in problems:
        doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('3.3 거대한 컴포넌트 - 유지보수 불가능 🔴', 2)
    
    component_data = [
        ['파일명', '라인 수', '상태'],
        ['PersonalServiceDashboard.tsx', '5,345', '⚠️⚠️⚠️ 위험'],
        ['App.tsx', '2,308', '⚠️ 과다'],
        ['AIResearchGuidePage.tsx', '2,154', '⚠️ 과다'],
        ['FuzzyAHPMethodologyPage.tsx', '1,511', '주의'],
        ['ComprehensiveUserGuide.tsx', '1,390', '주의']
    ]
    add_table_with_style(doc, component_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, 'PersonalServiceDashboard.tsx 분석:', bold=True)
    dashboard_issues = [
        '27개의 useState 훅 사용 (상태 관리 복잡도 극대화)',
        '단일 책임 원칙(SRP) 완전 위반',
        '테스트 불가능한 구조',
        '버그 수정 시 연쇄 오류 발생 위험',
        '새 기능 추가 시 복잡도 기하급수 증가',
        '팀 협업 불가능',
        '코드 리뷰 비현실적'
    ]
    for issue in dashboard_issues:
        para = doc.add_paragraph(issue, style='List Bullet')
        if '27개' in issue or 'SRP' in issue:
            for run in para.runs:
                run.font.color.rgb = RGBColor(220, 38, 38)
    
    doc.add_page_break()
    
    # ==================== 4. 중요 문제 ====================
    add_heading_with_color(doc, '4. 중요 문제 (High Priority Issues)', 1, RGBColor(249, 115, 22))
    
    doc.add_heading('4.1 타입 안정성 부족 🟠', 2)
    
    type_data = [
        ['지표', '현황'],
        ["타입 정의에 'any' 사용", '40회'],
        ["Props에 'any' 타입", '다수'],
        ['암묵적 any 허용', '예']
    ]
    add_table_with_style(doc, type_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '예시 (PersonalServiceDashboard.tsx):', bold=True)
    doc.add_paragraph('projects?: any[];')
    doc.add_paragraph('onCreateProject?: (projectData: any) => Promise<any>;')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '영향:', bold=True)
    doc.add_paragraph('런타임 에러 위험 증가', style='List Bullet')
    doc.add_paragraph('IDE 자동완성 불완전', style='List Bullet')
    doc.add_paragraph('리팩토링 시 오류 감지 불가', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('4.2 과도한 디버그 코드 🟠', 2)
    
    debug_data = [
        ['항목', '개수', '위험도'],
        ['console.log', '568개', '⚠️ 높음'],
        ['TODO/FIXME', '52개', '⚠️ 중간'],
        ['미완성 코드 마커', '52개', '주의']
    ]
    add_table_with_style(doc, debug_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('프로덕션 빌드에 디버그 코드 포함', style='List Bullet')
    doc.add_paragraph('콘솔 성능 저하', style='List Bullet')
    doc.add_paragraph('민감 정보 노출 위험', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('4.3 패키지 업데이트 지연 🟠', 2)
    
    outdated_data = [
        ['패키지', '현재 버전', '최신 버전', '격차'],
        ['TypeScript', '4.9.5', '5.9.3', '메이저 업데이트'],
        ['React', '18.3.1', '19.2.4', '메이저 업데이트'],
        ['React Router', '6.30.1', '7.13.0', '메이저 업데이트'],
        ['@anthropic-ai/claude-agent-sdk', '0.1.14', '0.2.44', '메이저 업데이트'],
        ['axios', '1.12.2', '1.13.5', '마이너 업데이트'],
        ['recharts', '2.15.4', '3.7.0', '메이저 업데이트']
    ]
    add_table_with_style(doc, outdated_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '위험도:', bold=True, color=RGBColor(249, 115, 22))
    doc.add_paragraph('보안 패치 누락', style='List Bullet')
    doc.add_paragraph('새 기능 미활용', style='List Bullet')
    doc.add_paragraph('호환성 문제 누적', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('4.4 테스트 커버리지 부족 🟠', 2)
    
    test_data = [
        ['항목', '현황'],
        ['테스트 파일', '20개'],
        ['컴포넌트 파일', '247개'],
        ['추정 커버리지', '< 10%']
    ]
    add_table_with_style(doc, test_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('회귀 테스트 불가능', style='List Bullet')
    doc.add_paragraph('리팩토링 위험 높음', style='List Bullet')
    doc.add_paragraph('CI/CD 파이프라인 불완전', style='List Bullet')
    doc.add_paragraph('품질 보증 불가', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 5. 경미한 문제 ====================
    add_heading_with_color(doc, '5. 경미한 문제 (Medium Priority Issues)', 1, RGBColor(234, 179, 8))
    
    doc.add_heading('5.1 ESLint 경고 다수 🟡', 2)
    
    add_paragraph_with_style(doc, '주요 경고 패턴:', bold=True)
    eslint_warnings = [
        'React Hook 의존성 배열 누락: 다수',
        '사용하지 않는 변수: 다수',
        '사용하지 않는 import: 다수',
        'no-loop-func 위반',
        'exhaustive-deps 경고'
    ]
    for warning in eslint_warnings:
        doc.add_paragraph(warning, style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('5.2 코드 중복 및 일관성 부족 🟡', 2)
    
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('유사한 컴포넌트 패턴 반복', style='List Bullet')
    doc.add_paragraph('DRY 원칙 위반', style='List Bullet')
    doc.add_paragraph('네이밍 컨벤션 불일치 (예: PersonalServiceDashboard vs PersonalServiceDashboard_Enhanced)', style='List Bullet')
    doc.add_paragraph('중복 서비스 파일 (dataService.ts vs dataService_clean.ts)', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('5.3 문서화 과잉 vs 코드 주석 부족 🟡', 2)
    
    doc_data = [
        ['항목', '현황'],
        ['마크다운 문서 (Dev_md_2/)', '92개'],
        ['코드 내 주석', '부족'],
        ['API 문서', '미비'],
        ['컴포넌트 사용법 설명', '없음']
    ]
    add_table_with_style(doc, doc_data)
    
    doc.add_page_break()
    
    # ==================== 6. 성능 문제 분석 ====================
    add_heading_with_color(doc, '6. 성능 문제 분석', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('6.1 번들 크기 상세 분석', 2)
    
    add_paragraph_with_style(doc, '현재 번들 구성 (추정):', bold=True)
    
    bundle_breakdown = [
        ['컴포넌트', '예상 크기'],
        ['React + React-DOM', '~140KB'],
        ['Recharts', '~500KB'],
        ['FontAwesome', '~300KB'],
        ['Tailwind CSS', '~50KB (최적화 후)'],
        ['애플리케이션 코드', '~1.8MB ⚠️'],
        ['합계', '2.8MB']
    ]
    add_table_with_style(doc, bundle_breakdown)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True, color=RGBColor(220, 38, 38))
    doc.add_paragraph('모든 컴포넌트가 메인 번들에 포함', style='List Bullet')
    doc.add_paragraph('라우트별 분리 안됨', style='List Bullet')
    doc.add_paragraph('초기 렌더링에 불필요한 코드 로드', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('6.2 개선 목표', 2)
    
    improvement_data = [
        ['항목', '현재', '목표'],
        ['초기 번들 크기', '2.8MB', '< 500KB'],
        ['초기 로딩 시간', '~5초', '< 1.5초'],
        ['Time to Interactive', '~8초', '< 3초']
    ]
    add_table_with_style(doc, improvement_data)
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '개선 방법:', bold=True)
    doc.add_paragraph('React.lazy() + Suspense 적용', style='List Bullet')
    doc.add_paragraph('라우트 기반 코드 스플리팅', style='List Bullet')
    doc.add_paragraph('동적 import 활용', style='List Bullet')
    doc.add_paragraph('트리 쉐이킹 최적화', style='List Bullet')
    doc.add_paragraph('이미지 및 에셋 최적화', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 7. 아키텍처 문제 ====================
    add_heading_with_color(doc, '7. 아키텍처 문제', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('7.1 상태 관리 부재', 2)
    
    add_paragraph_with_style(doc, '현재 상황:', bold=True)
    doc.add_paragraph('Props Drilling 방식 사용', style='List Bullet')
    doc.add_paragraph('전역 상태 관리 라이브러리 없음', style='List Bullet')
    doc.add_paragraph('깊은 컴포넌트 트리에서 props 전달 복잡', style='List Bullet')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '권장 솔루션:', bold=True, color=RGBColor(34, 197, 94))
    doc.add_paragraph('Redux Toolkit 도입', style='List Bullet')
    doc.add_paragraph('또는 Zustand (경량 대안)', style='List Bullet')
    doc.add_paragraph('React Query (서버 상태 관리)', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('7.2 컴포넌트 계층 구조 문제', 2)
    
    add_paragraph_with_style(doc, '현재 구조 (문제):', bold=True, color=RGBColor(220, 38, 38))
    doc.add_paragraph('PersonalServiceDashboard (5,345 라인)')
    doc.add_paragraph('  ├── 27개 useState 훅', style='List Bullet 2')
    doc.add_paragraph('  ├── 수십 개 자식 컴포넌트 직접 관리', style='List Bullet 2')
    doc.add_paragraph('  └── 비즈니스 로직 + UI 로직 혼재', style='List Bullet 2')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '권장 구조:', bold=True, color=RGBColor(34, 197, 94))
    doc.add_paragraph('PersonalServiceDashboard (< 200 라인)')
    doc.add_paragraph('  ├── ProjectSection (Container)', style='List Bullet 2')
    doc.add_paragraph('  │   ├── ProjectList (Presentation)', style='List Bullet 3')
    doc.add_paragraph('  │   └── ProjectForm (Presentation)', style='List Bullet 3')
    doc.add_paragraph('  ├── CriteriaSection (Container)', style='List Bullet 2')
    doc.add_paragraph('  └── EvaluationSection (Container)', style='List Bullet 2')
    
    doc.add_paragraph('')
    doc.add_heading('7.3 API 레이어 분산', 2)
    
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('22개 서비스 파일 존재', style='List Bullet')
    doc.add_paragraph('역할 분담 불명확', style='List Bullet')
    doc.add_paragraph('중복 로직 다수', style='List Bullet')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '서비스 파일 중복 예시:', bold=True)
    doc.add_paragraph('api.ts vs apiService.ts (역할 불명확)', style='List Bullet')
    doc.add_paragraph('dataService.ts vs dataService_clean.ts (왜 2개?)', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 8. 데이터베이스 및 백엔드 문제 ====================
    add_heading_with_color(doc, '8. 데이터베이스 및 백엔드 문제', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('8.1 백엔드 의존성', 2)
    
    add_paragraph_with_style(doc, '프로덕션 백엔드:', bold=True)
    doc.add_paragraph('Render.com (유료 플랜)', style='List Bullet')
    doc.add_paragraph('URL: https://ahp-django-backend.onrender.com', style='List Bullet')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '위험 요소:', bold=True, color=RGBColor(220, 38, 38))
    doc.add_paragraph('무료 플랜 종료 시 서비스 중단 위험', style='List Bullet')
    doc.add_paragraph('백업 전략 불명확', style='List Bullet')
    doc.add_paragraph('장애 복구 계획 없음', style='List Bullet')
    doc.add_paragraph('CORS 문제 잠재적 위험', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('8.2 오프라인 지원 부재', 2)
    
    add_paragraph_with_style(doc, '현재 설정:', bold=True)
    doc.add_paragraph('REACT_APP_DATA_MODE=online')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('인터넷 끊김 시 동작 불가', style='List Bullet')
    doc.add_paragraph('Progressive Web App (PWA) 미구현', style='List Bullet')
    doc.add_paragraph('로컬 캐싱 전략 없음', style='List Bullet')
    doc.add_paragraph('Service Worker 미사용', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 9. 개발 환경 문제 ====================
    add_heading_with_color(doc, '9. 개발 환경 문제', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('9.1 CI/CD 파이프라인 불완전', 2)
    
    add_paragraph_with_style(doc, '현재 상황:', bold=True)
    doc.add_paragraph('GitHub Actions 사용 중', style='List Bullet')
    doc.add_paragraph('deploy.yml 존재', style='List Bullet')
    
    doc.add_paragraph('')
    add_paragraph_with_style(doc, '문제점:', bold=True)
    doc.add_paragraph('테스트 단계 스킵 가능 (skip_tests 옵션)', style='List Bullet')
    doc.add_paragraph('빌드 실패 시 롤백 전략 없음', style='List Bullet')
    doc.add_paragraph('환경별 배포 전략 없음', style='List Bullet')
    doc.add_paragraph('스테이징 환경 미구축', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('9.2 환경 설정 문제', 2)
    
    add_paragraph_with_style(doc, '발견된 문제:', bold=True, color=RGBColor(220, 38, 38))
    doc.add_paragraph('.env 파일이 git에 포함 (보안 위험)', style='List Bullet')
    doc.add_paragraph('환경 변수 검증 로직 없음', style='List Bullet')
    doc.add_paragraph('로컬/스테이징/프로덕션 환경 구분 불명확', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 10. 우선순위별 해결 로드맵 ====================
    add_heading_with_color(doc, '10. 우선순위별 해결 로드맵', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('10.1 즉시 조치 (1주 이내) 🚨', 2)
    
    immediate_tasks = [
        '보안 취약점 패치 (npm audit fix)',
        'console.log 568개 제거 (프로덕션 빌드)',
        '.env 파일 gitignore 추가 및 환경 변수 관리',
        '타입 안정성 개선 (any → 명시적 타입)',
        'ESLint 경고 50개 이하로 감소'
    ]
    
    for i, task in enumerate(immediate_tasks, 1):
        para = doc.add_paragraph(f'{i}. {task}')
        for run in para.runs:
            if '보안' in task or 'console.log' in task:
                run.font.color.rgb = RGBColor(220, 38, 38)
                run.font.bold = True
    
    doc.add_paragraph('')
    doc.add_heading('10.2 단기 개선 (1개월 이내) 🔥', 2)
    
    short_term_tasks = [
        '번들 크기 최적화 (코드 스플리팅, 목표: < 500KB)',
        'PersonalServiceDashboard.tsx 리팩토링 (목표: < 500 라인)',
        '패키지 업데이트 (TypeScript 5.x, React 19.x)',
        '테스트 커버리지 50% 달성',
        'PWA 기본 구현 (Service Worker, Manifest)',
        '성능 모니터링 도구 통합 (Lighthouse CI)'
    ]
    
    for i, task in enumerate(short_term_tasks, 1):
        doc.add_paragraph(f'{i}. {task}')
    
    doc.add_paragraph('')
    doc.add_heading('10.3 중기 개선 (3개월 이내) 🎯', 2)
    
    mid_term_tasks = [
        '상태 관리 라이브러리 도입 (Redux Toolkit 또는 Zustand)',
        '아키텍처 재설계 (레이어 분리, Container/Presentation 패턴)',
        'API 레이어 통합 및 정리',
        '컴포넌트 라이브러리 구축 (Storybook)',
        'E2E 테스트 구축 (Playwright 또는 Cypress)',
        '스테이징 환경 구축',
        '백업 및 복구 전략 수립'
    ]
    
    for i, task in enumerate(mid_term_tasks, 1):
        doc.add_paragraph(f'{i}. {task}')
    
    doc.add_paragraph('')
    doc.add_heading('10.4 장기 개선 (6개월 이내) 📈', 2)
    
    long_term_tasks = [
        '마이크로 프론트엔드 아키텍처 검토',
        '국제화(i18n) 완전 구현',
        '접근성(a11y) 개선 (WCAG 2.1 AA 준수)',
        '고급 분석 기능 완성 (Monte Carlo 시뮬레이션)',
        '모바일 앱 개발 (React Native)',
        '결제 시스템 구현',
        'AI 기능 고도화'
    ]
    
    for i, task in enumerate(long_term_tasks, 1):
        doc.add_paragraph(f'{i}. {task}')
    
    doc.add_page_break()
    
    # ==================== 11. 긍정적 측면 ====================
    add_heading_with_color(doc, '11. 긍정적 측면 (칭찬할 점)', 1, RGBColor(34, 197, 94))
    
    add_paragraph_with_style(doc, '잘 된 부분:', bold=True, size=12)
    
    doc.add_paragraph('')
    
    positive_aspects = [
        ('TypeScript 사용', 'TypeScript 기반으로 타입 안정성 기본 확보'),
        ('Tailwind CSS', '일관된 디자인 시스템 구축'),
        ('명확한 프로젝트 구조', 'components 디렉토리 체계적 분류'),
        ('풍부한 문서화', '92개 개발 문서로 프로젝트 히스토리 완벽 보존'),
        ('CI/CD 구축', 'GitHub Actions 활용한 자동 배포'),
        ('모던 스택', 'React 18, Django 5.0 등 최신 기술 스택'),
        ('@ts-ignore 0개', 'TypeScript 규칙을 우회하지 않는 정직한 코드'),
        ('실제 백엔드 연동', 'PostgreSQL 데이터베이스 실제 사용'),
        ('반응형 디자인', '모바일/태블릿/데스크톱 지원'),
        ('계층적 평가 시스템', 'Power Method 알고리즘 완전 구현')
    ]
    
    for title, desc in positive_aspects:
        para = doc.add_paragraph()
        run = para.add_run(f'✅ {title}: ')
        run.font.bold = True
        run.font.color.rgb = RGBColor(34, 197, 94)
        run = para.add_run(desc)
    
    doc.add_page_break()
    
    # ==================== 12. 최종 권고사항 ====================
    add_heading_with_color(doc, '12. 최종 권고사항', 1, RGBColor(37, 99, 235))
    
    doc.add_heading('12.1 즉시 실행해야 할 명령어', 2)
    
    doc.add_paragraph('# 1. 보안 패치', style='Heading 3')
    code_para = doc.add_paragraph('npm audit fix')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    code_para = doc.add_paragraph('npm audit fix --force  # 주의: breaking changes 가능')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    doc.add_paragraph('')
    doc.add_paragraph('# 2. 디버그 코드 제거', style='Heading 3')
    code_para = doc.add_paragraph('grep -r "console.log" src/ --files-with-matches')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    doc.add_paragraph('')
    doc.add_paragraph('# 3. 불필요한 의존성 제거', style='Heading 3')
    code_para = doc.add_paragraph('npm prune')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    code_para = doc.add_paragraph('npx depcheck')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    doc.add_paragraph('')
    doc.add_paragraph('# 4. 빌드 분석', style='Heading 3')
    code_para = doc.add_paragraph('npm install --save-dev webpack-bundle-analyzer')
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    doc.add_paragraph('')
    doc.add_heading('12.2 리팩토링 전략', 2)
    
    refactoring_phases = [
        ('1단계', '긴급 수정 (보안, 성능)', '1주'),
        ('2단계', '컴포넌트 분리 (PersonalServiceDashboard)', '2주'),
        ('3단계', '상태 관리 도입', '2주'),
        ('4단계', '테스트 작성', '3주'),
        ('5단계', '문서화 업데이트', '1주')
    ]
    
    refactoring_table = [['단계', '작업 내용', '예상 기간']]
    refactoring_table.extend(refactoring_phases)
    add_table_with_style(doc, refactoring_table)
    
    doc.add_page_break()
    
    # ==================== 13. 결론 ====================
    add_heading_with_color(doc, '13. 결론', 1, RGBColor(37, 99, 235))
    
    doc.add_paragraph('')
    
    para = doc.add_paragraph()
    run = para.add_run('당신의 AHP 플랫폼은 ')
    run.font.size = Pt(12)
    run = para.add_run('기능적으로는 훌륭하지만')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(34, 197, 94)
    run = para.add_run(', ')
    run.font.size = Pt(12)
    run = para.add_run('프로덕션 배포에는 아직 준비되지 않았습니다')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = para.add_run('.')
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    doc.add_heading('핵심 문제 3가지', 2)
    
    core_issues = [
        ('🔴 보안 취약점', '즉시 패치 필요 (React Router XSS, ajv ReDoS)'),
        ('🔴 2.8MB 번들', '성능 재앙 수준 - 초기 로딩 5초 이상'),
        ('🔴 5,345 라인 컴포넌트', '유지보수 불가능 - 리팩토링 필수')
    ]
    
    for issue, desc in core_issues:
        para = doc.add_paragraph()
        run = para.add_run(issue + ': ')
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(220, 38, 38)
        run = para.add_run(desc)
        run.font.size = Pt(11)
    
    doc.add_paragraph('')
    doc.add_heading('하지만 희망적인 점', 2)
    
    hope_points = [
        '✅ 탄탄한 기능 구현 (AHP 알고리즘 완전 구현)',
        '✅ 명확한 비전 (연구자 특화 플랫폼)',
        '✅ 체계적인 문서화 (92개 개발 문서)',
        '✅ 현대적 기술 스택 (React 18, TypeScript, Django 5.0)',
        '✅ 실제 사용자 피드백 반영'
    ]
    
    for point in hope_points:
        para = doc.add_paragraph(point)
        for run in para.runs:
            run.font.color.rgb = RGBColor(34, 197, 94)
            run.font.size = Pt(11)
    
    doc.add_paragraph('')
    
    para = doc.add_paragraph()
    run = para.add_run('1~2개월의 집중적인 리팩토링')
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(37, 99, 235)
    run = para.add_run('으로 ')
    run.font.size = Pt(13)
    run = para.add_run('프로덕션 레벨')
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(34, 197, 94)
    run = para.add_run('로 끌어올릴 수 있습니다! 💪')
    run.font.size = Pt(13)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 마무리
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run('--- 보고서 끝 ---')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(107, 114, 128)
    
    # 파일 저장
    filename = f'/home/user/webapp/AHP_플랫폼_분석보고서_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f'✅ 보고서 생성 완료: {filename}')
    return filename

if __name__ == '__main__':
    create_analysis_report()
